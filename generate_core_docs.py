"""
generate_core_docs.py — Build a comprehensive Word document for all core
paper-trading scripts.

Run:  python3 generate_core_docs.py
Output: Documentation/doc_core_paper_trading_system.docx
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    if level > 0:
        p.paragraph_format.left_indent = Pt(18 * level)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    return p


def add_table_row(table, cells_text):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        row.cells[i].text = str(text)
    return row


def build_document():
    doc = Document()

    # ─── Title Page ───────────────────────────────────────────────────────
    title = doc.add_heading("Stock Market AI Bot", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading("Core Paper Trading System Documentation", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "A complete guide for beginners with no prior experience", italic=True).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Generated: May 2026", italic=True).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ─── Table of Contents (manual) ───────────────────────────────────────
    add_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. System Overview",
        "2. How the System Works (End-to-End)",
        "3. core_satellite_alpha.py - Moomoo Signal Generator",
        "4. core_satellite_tqqq.py - Alpaca TQQQ Signal Generator",
        "5. moomoo_paper_trading.py - Moomoo Broker Script",
        "6. alpaca_paper_trading.py - Alpaca Broker Script",
        "7. paper_gauntlet.py - Moomoo Health Check",
        "8. alpaca_paper_gauntlet.py - Alpaca Health Check",
        "9. daily_run.py - Automation Runner",
        "10. paper_report.py - Performance Comparison",
        "11. settings.py - Configuration",
        "12. Glossary of Key Terms",
    ]
    for item in toc_items:
        add_para(doc, item)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 1: SYSTEM OVERVIEW
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "1. System Overview", level=1)

    add_para(doc, "What is this system?", bold=True)
    add_para(doc, (
        "This is an automated stock trading system that uses machine learning and "
        "quantitative factors to decide what stocks and ETFs to buy. Instead of "
        "trading with real money right away, it first runs on 'paper trading' accounts "
        "- fake-money accounts that simulate real trading. This lets you prove the "
        "strategy works before risking actual capital."
    ))

    add_para(doc, "The Two Strategies", bold=True)
    add_para(doc, (
        "The system runs two strategies side-by-side on two different brokers:"
    ))

    # Strategy comparison table
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Feature"
    hdr[1].text = "Moomoo (Core-Satellite)"
    hdr[2].text = "Alpaca (TQQQ-Enhanced)"

    rows_data = [
        ("Broker", "Moomoo (via OpenD desktop app)", "Alpaca (free cloud API)"),
        ("Signal File", "core_satellite_alpha_signal.csv", "core_satellite_tqqq_signal.csv"),
        ("Core Holdings", "SPY + QQQ", "SPY + QQQ + TQQQ"),
        ("TQQQ (3x Leverage)", "No", "Yes, during uptrends only"),
        ("Overlay Stocks", "Top 3 by factor score", "Top 3 by factor score"),
        ("Backtested Return", "~8,887%", "~10,369%"),
        ("Max Drawdown Limit", "-10%", "-15% (wider for TQQQ)"),
    ]
    for row_data in rows_data:
        add_table_row(table, row_data)

    add_para(doc, "")
    add_para(doc, "What is 'Core-Satellite'?", bold=True)
    add_para(doc, (
        "Core-Satellite is an investment approach where you split your money into two parts:\n\n"
        "- The 'Core' (50-70% of money): Safe, diversified ETFs like SPY (S&P 500) and QQQ (Nasdaq 100). "
        "These give you broad market exposure.\n\n"
        "- The 'Satellite' (30-50% of money): Individual stock picks chosen by the AI's factor scoring system. "
        "These are the stocks the system believes will outperform the market."
    ))

    add_para(doc, "What is Regime Switching?", bold=True)
    add_para(doc, (
        "The system doesn't use the same allocation all the time. It detects the current "
        "'market regime' - whether the market is bullish, neutral, or bearish - and adjusts accordingly:\n\n"
        "- Risk On: Both SPY and QQQ are above their moving averages AND volatility is low. "
        "The system goes aggressive: more QQQ, more stock picks, and (for Alpaca) adds TQQQ.\n\n"
        "- Neutral: SPY is doing fine but QQQ isn't, or volatility is rising. The system goes balanced.\n\n"
        "- Risk Off: Markets are falling. The system gets defensive: more SPY, fewer stock picks, zero TQQQ."
    ))

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 2: END-TO-END WORKFLOW
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "2. How the System Works (End-to-End)", level=1)

    add_para(doc, "The Daily Workflow", bold=True)
    add_para(doc, (
        "Every trading day, the system runs these steps in order. You can run them "
        "manually one by one, or use daily_run.py to automate all of them:"
    ))

    steps = [
        ("Step 1: Generate Moomoo Signal", "core_satellite_alpha.py",
         "Analyzes factor data, detects the market regime, picks top stocks, and writes "
         "a signal CSV file telling Moomoo what to buy/sell."),
        ("Step 2: Submit Moomoo Orders", "moomoo_paper_trading.py --submit",
         "Reads the signal CSV, compares it to your current Moomoo positions, generates "
         "buy/sell orders, and submits them to Moomoo's paper trading account."),
        ("Step 3: Sync Moomoo Status", "moomoo_paper_trading.py --status",
         "Connects to Moomoo to get your current account equity and positions, saves "
         "a daily snapshot for performance tracking."),
        ("Step 4: Moomoo Health Check", "paper_gauntlet.py",
         "Evaluates whether the Moomoo strategy is performing well enough. Checks Sharpe ratio, "
         "drawdown, fill rate, and other safety metrics."),
        ("Step 5: Generate Alpaca Signal", "core_satellite_tqqq.py",
         "Same as Step 1 but generates a signal that includes TQQQ allocation during uptrends."),
        ("Step 6: Submit Alpaca Orders", "alpaca_paper_trading.py --submit",
         "Reads the TQQQ signal, compares to current Alpaca positions, submits orders. "
         "Also automatically saves an equity snapshot."),
        ("Step 7: Reconcile Alpaca Orders", "alpaca_paper_trading.py --reconcile",
         "Checks with Alpaca's API to see if submitted orders actually filled, and updates "
         "the trade log with real fill statuses."),
        ("Step 8: Alpaca Health Check", "alpaca_paper_gauntlet.py",
         "Same as Step 4 but for Alpaca. Uses wider drawdown limits because TQQQ is more volatile."),
    ]

    for title, cmd, desc in steps:
        add_para(doc, title, bold=True)
        add_code(doc, f"  Command: python3 {cmd}")
        add_para(doc, desc)

    add_para(doc, "")
    add_para(doc, "The Flow Diagram", bold=True)
    add_code(doc, (
        "  Factor Data --> core_satellite_alpha.py --> signal CSV --> moomoo_paper_trading.py --> Moomoo Broker\n"
        "                                                                                              |\n"
        "                                                                                    paper_gauntlet.py\n"
        "                                                                                              |\n"
        "  Factor Data --> core_satellite_tqqq.py ---> signal CSV --> alpaca_paper_trading.py --> Alpaca Broker\n"
        "                                                                                              |\n"
        "                                                                                  alpaca_paper_gauntlet.py\n"
        "                                                                                              |\n"
        "                                                                                     paper_report.py\n"
        "                                                                                    (compare both)"
    ))

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 3: core_satellite_alpha.py
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "3. core_satellite_alpha.py - Moomoo Signal Generator", level=1)

    add_para(doc, "Purpose", bold=True)
    add_para(doc, (
        "This script is the brain of the Moomoo strategy. It analyzes historical stock data, "
        "runs a grid search over hundreds of configuration combinations, picks the best one, "
        "and writes a signal file that tells the broker script exactly what to buy and sell."
    ))

    add_para(doc, "How to Run", bold=True)
    add_code(doc, "  python3 core_satellite_alpha.py")
    add_para(doc, (
        "No flags needed. It runs the full grid search and generates the signal automatically. "
        "Takes several minutes because it tests hundreds of configurations."
    ))

    add_para(doc, "What It Outputs", bold=True)
    add_bullet(doc, "signals/core_satellite_alpha_signal.csv - The trading signal for Moomoo")
    add_bullet(doc, "signals/core_satellite_alpha_grid.csv - Results of every config tested")
    add_bullet(doc, "signals/core_satellite_alpha_metrics.json - Detailed performance metrics")
    add_bullet(doc, "signals/core_satellite_alpha_equity.csv - Backtested equity curve")
    add_bullet(doc, "signals/core_satellite_alpha_trades.csv - Every trade in the backtest")

    add_para(doc, "Key Logic Explained", bold=True)

    add_para(doc, "1. Regime Detection (_resolve_allocation)", italic=True)
    add_para(doc, (
        "The system checks three things every rebalance day to determine the market regime:\n\n"
        "- Is SPY above its 200-day moving average? (Long-term uptrend check)\n"
        "- Is QQQ above its 100-day moving average? (Tech sector health check)\n"
        "- Is QQQ's 20-day realized volatility below 30%? (Calm markets check)\n\n"
        "If all three are YES --> risk_on (aggressive allocation)\n"
        "If SPY is OK but QQQ fails or vol is high --> neutral (balanced allocation)\n"
        "If SPY fails --> risk_off (defensive allocation)\n\n"
        "This is the most important decision the system makes. It controls how much "
        "money goes into stocks vs safe ETFs."
    ))

    add_para(doc, "2. Stock Selection (_select_sticky_holdings)", italic=True)
    add_para(doc, (
        "Once the regime is known, the system picks the top 3 stocks from its universe "
        "using 'factor scores'. These scores measure how attractive each stock is based on "
        "things like momentum, value, and quality.\n\n"
        "'Sticky' means the system prefers to keep stocks it already holds rather than constantly "
        "swapping them. A held stock only gets dropped if its score rank falls below the "
        "'exit rank floor' (default: 80th percentile). This reduces trading costs."
    ))

    add_para(doc, "3. Weight Assignment (_overlay_weights / _sticky_overlay_weights)", italic=True)
    add_para(doc, (
        "After picking the stocks, the system decides how much money to put in each one. "
        "Higher-scored stocks get more weight. No single stock can exceed 35% of the overlay "
        "allocation (the MAX_SINGLE_NAME_WEIGHT cap).\n\n"
        "'Sticky scoring' means new weights are blended 65% with the previous period's weights. "
        "This prevents dramatic shifts in position sizes from one rebalance to the next."
    ))

    add_para(doc, "4. Paper Safety Scaling (_scale_paper_targets_to_gross)", italic=True)
    add_para(doc, (
        "The backtest might test 1.25x gross exposure (investing 125% of your money using leverage), "
        "but paper trading accounts usually reject leverage. This function scales everything down "
        "proportionally so total exposure stays at or below 1.00x (100% of your money)."
    ))

    add_para(doc, "5. Grid Search (main function)", italic=True)
    add_para(doc, (
        "The script tests hundreds of configuration combinations:\n"
        "- 2 core presets (40/60 SPY/QQQ vs 25/75)\n"
        "- 5 regime presets (different moving average windows and overlay sizes)\n"
        "- 4 score sources (different factor scoring methods)\n"
        "- 3 cost stress levels (2x, 3x, 5x transaction costs)\n"
        "- 2 holding periods (10 days vs 20 days)\n"
        "- And more...\n\n"
        "It then picks the best configuration that passes all quality gates. "
        "A configuration must beat SPY, QQQ, and a 60/40 blend across multiple time periods "
        "and survive cost stress tests before being used for live trading."
    ))

    add_para(doc, "Key Parameters", bold=True)
    param_table = doc.add_table(rows=1, cols=3)
    param_table.style = "Light Grid Accent 1"
    hdr = param_table.rows[0].cells
    hdr[0].text = "Parameter"
    hdr[1].text = "Default"
    hdr[2].text = "What It Means"

    params = [
        ("PAPER_MAX_GROSS_EXPOSURE", "1.00", "Maximum total investment as fraction of account (1.0 = 100%, no leverage)"),
        ("MAX_SINGLE_NAME_WEIGHT", "0.35", "No single stock can be more than 35% of the overlay"),
        ("MAX_POSITIVE_YEAR_ALPHA_SHARE", "0.35", "No single year can contribute more than 35% of total alpha (prevents overfitting to one lucky year)"),
        ("PAPER_SIGNAL_TIMEZONE", "Asia/Singapore", "Timezone used when stamping signal generation time"),
        ("CORE_OVERLAY_COMBOS", "(1.0/0.25), (0.75/0.25), (0.75/0.50)", "Core/overlay gross exposure ratios tested in grid search"),
        ("HOLDING_DAY_OPTIONS", "10, 20", "How many trading days between rebalances"),
        ("EXIT_RANK_FLOORS", "0.80", "Held stocks must stay above 80th percentile score or get dropped"),
        ("COST_STRESS_MULTIPLIERS", "2.0, 3.0, 5.0", "Transaction cost multipliers used in stress testing"),
    ]
    for p in params:
        add_table_row(param_table, p)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 4: core_satellite_tqqq.py
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "4. core_satellite_tqqq.py - Alpaca TQQQ Signal Generator", level=1)

    add_para(doc, "Purpose", bold=True)
    add_para(doc, (
        "This script generates a trading signal for the Alpaca paper account. It's identical "
        "to the Moomoo signal generator except for one key difference: during risk_on periods "
        "(confirmed uptrends with low volatility), it replaces 20% of the QQQ core allocation "
        "with TQQQ - a 3x leveraged ETF that triples QQQ's daily returns."
    ))

    add_para(doc, "How to Run", bold=True)
    add_code(doc, "  python3 core_satellite_tqqq.py                     # Generate live signal (DEFAULT)")
    add_code(doc, "  python3 core_satellite_tqqq.py --tqqq-weight 0.30  # Use 30% TQQQ instead of 20%")
    add_code(doc, "  python3 core_satellite_tqqq.py --backtest          # Run historical backtest")
    add_code(doc, "  python3 core_satellite_tqqq.py --grid              # Grid search over TQQQ weights")

    add_para(doc, "")
    add_para(doc, "What is TQQQ?", bold=True)
    add_para(doc, (
        "TQQQ is a leveraged ETF that aims to deliver 3 times QQQ's daily return. "
        "If QQQ goes up 1% today, TQQQ should go up about 3%. But if QQQ drops 1%, "
        "TQQQ drops about 3%. It amplifies both gains AND losses.\n\n"
        "This is why the system ONLY holds TQQQ during risk_on periods. During neutral "
        "or risk_off, TQQQ weight is forced to zero."
    ))

    add_para(doc, "Key Logic: TQQQ Allocation by Regime", bold=True)
    regime_table = doc.add_table(rows=1, cols=5)
    regime_table.style = "Light Grid Accent 1"
    hdr = regime_table.rows[0].cells
    hdr[0].text = "Regime"
    hdr[1].text = "SPY Weight"
    hdr[2].text = "QQQ Weight"
    hdr[3].text = "TQQQ Weight"
    hdr[4].text = "Overlay"

    regime_rows = [
        ("risk_on", "0%", "80% of core", "20% of core", "70% gross"),
        ("neutral", "25%", "75%", "0% (too risky)", "70% gross"),
        ("risk_off", "60%", "40%", "0% (never)", "35% gross"),
    ]
    for r in regime_rows:
        add_table_row(regime_table, r)

    add_para(doc, "")
    add_para(doc, "Key Logic Explained", bold=True)

    add_para(doc, "1. build_tqqq_presets(tqqq_weight)", italic=True)
    add_para(doc, (
        "Creates the regime allocation rules. The tqqq_weight parameter (default 0.20) "
        "controls what fraction of the risk_on core goes to TQQQ. Example: with "
        "tqqq_weight=0.20 and core_gross=0.55, the risk_on core allocates:\n"
        "  - QQQ: 0.55 * 0.80 = 0.44 (44% of portfolio)\n"
        "  - TQQQ: 0.55 * 0.20 = 0.11 (11% of portfolio)"
    ))

    add_para(doc, "2. write_tqqq_signal(panel, tqqq_weight)", italic=True)
    add_para(doc, (
        "The main signal generation function. It:\n"
        "a) Looks at the most recent factor data to determine today's date\n"
        "b) Detects the current market regime (risk_on/neutral/risk_off)\n"
        "c) Picks the top 3 stocks using factor scores appropriate for the regime\n"
        "d) Assigns weights to the stocks\n"
        "e) Computes ETF target weights (SPY, QQQ, TQQQ) based on regime\n"
        "f) Scales everything down to stay within 1.0x gross exposure\n"
        "g) Writes the signal CSV to signals/core_satellite_tqqq_signal.csv"
    ))

    add_para(doc, "3. _scale_paper_targets()", italic=True)
    add_para(doc, (
        "If the strategy says to invest 125% of your money (core_gross=0.55 + overlay_gross=0.70 = 1.25x), "
        "this function scales everything down proportionally to fit within the "
        "PAPER_MAX_GROSS_EXPOSURE limit (default 1.00). Every weight gets multiplied by "
        "the same scale factor, so the relative proportions stay the same."
    ))

    add_para(doc, "Command Line Parameters", bold=True)
    cli_table = doc.add_table(rows=1, cols=3)
    cli_table.style = "Light Grid Accent 1"
    hdr = cli_table.rows[0].cells
    hdr[0].text = "Parameter"
    hdr[1].text = "Default"
    hdr[2].text = "What It Does"

    cli_params = [
        ("--tqqq-weight", "0.20", "What fraction of risk_on core goes to TQQQ (0.20 = 20%)"),
        ("--backtest", "Off", "Run historical backtest instead of generating a live signal"),
        ("--grid", "Off", "Test multiple TQQQ weights (0%, 10%, 15%, 20%, 25%, 30%, 40%, 50%)"),
        ("--cost-stress", "2.0", "Multiply estimated transaction costs by this factor"),
        ("--holding-days", "10", "How many trading days between rebalances"),
    ]
    for p in cli_params:
        add_table_row(cli_table, p)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 5: moomoo_paper_trading.py
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "5. moomoo_paper_trading.py - Moomoo Broker Script", level=1)

    add_para(doc, "Purpose", bold=True)
    add_para(doc, (
        "This is the bridge between the strategy's signal and the Moomoo broker. It reads "
        "the signal CSV, connects to Moomoo's desktop app (OpenD), calculates what orders "
        "are needed, and submits them. It also tracks positions, reconciles fills, and saves "
        "daily equity snapshots."
    ))

    add_para(doc, "Prerequisites", bold=True)
    add_bullet(doc, "Moomoo OpenD desktop app must be running on your computer")
    add_bullet(doc, "pip install moomoo-api")
    add_bullet(doc, "Set PAPER_MODE_STRATEGY=core_satellite_alpha in your .env file")

    add_para(doc, "How to Run", bold=True)
    add_code(doc, "  python3 moomoo_paper_trading.py                     # Preview orders (dry run)")
    add_code(doc, "  python3 moomoo_paper_trading.py --submit            # Submit orders + auto-sync fills")
    add_code(doc, "  python3 moomoo_paper_trading.py --status            # Save equity/positions snapshot")
    add_code(doc, "  python3 moomoo_paper_trading.py --sync              # Reconcile fills from Moomoo history")
    add_code(doc, "  python3 moomoo_paper_trading.py --force --submit    # Override rebalance guard")

    add_para(doc, "Key Logic Explained", bold=True)

    add_para(doc, "1. Signal Loading (load_core_satellite_signal)", italic=True)
    add_para(doc, (
        "Reads core_satellite_alpha_signal.csv. Checks that paper_ready=True and "
        "gates_all_pass=True before allowing any trades. If either is False, the "
        "script refuses to trade - this is a safety mechanism."
    ))

    add_para(doc, "2. Weight Extraction (core_satellite_target_weights)", italic=True)
    add_para(doc, (
        "Reads the target_spy_weight, target_qqq_weight, and overlay_weights_json "
        "fields from the signal. Returns a dictionary like: "
        "{'SPY': 0.30, 'QQQ': 0.40, 'AAPL': 0.10, 'MSFT': 0.08, 'GOOGL': 0.07}"
    ))

    add_para(doc, "3. Drift-Based Rebalancing (add_drift_and_guards)", italic=True)
    add_para(doc, (
        "The system doesn't trade every tiny difference. It calculates the 'drift' - "
        "the gap between your current weight and target weight for each ticker. "
        "An order is only generated if:\n"
        "- ETF drift > 3% (e.g., SPY target is 30% but you hold 26%)\n"
        "- Stock drift > 1%\n"
        "- Or the position is being fully exited/entered\n\n"
        "This prevents expensive tiny trades that cost more than they're worth."
    ))

    add_para(doc, "4. Order Submission (submit_core_satellite_orders)", italic=True)
    add_para(doc, (
        "Connects to Moomoo via the local OpenD app, unlocks the trade context if needed, "
        "and submits limit orders. Sells go first (to free up cash), then buys. "
        "Each order uses a limit price with a small offset (default 10 basis points) - "
        "buys pay slightly above market price, sells go slightly below."
    ))

    add_para(doc, "5. Fill Reconciliation (reconcile_paper_trades)", italic=True)
    add_para(doc, (
        "After submitting orders, this checks Moomoo's order history to see if they "
        "actually filled. It matches each submitted order to a Moomoo order by order ID, "
        "ticker, side, and quantity. It also computes execution slippage - how much worse "
        "(or better) the fill price was compared to the planned price."
    ))

    add_para(doc, "6. Signal Freshness Validation (validate_signal_freshness)", italic=True)
    add_para(doc, (
        "Before trading, the script checks:\n"
        "- Is the signal less than 24 hours old? (configurable)\n"
        "- Is the underlying factor data less than 5 trading days old?\n\n"
        "If either check fails, trading is blocked to prevent acting on stale information."
    ))

    add_para(doc, "Key Parameters", bold=True)
    moomoo_params = doc.add_table(rows=1, cols=3)
    moomoo_params.style = "Light Grid Accent 1"
    hdr = moomoo_params.rows[0].cells
    hdr[0].text = "Parameter"
    hdr[1].text = "Default"
    hdr[2].text = "What It Does"

    mparams = [
        ("--submit", "Off", "Actually send orders to Moomoo. Without this, only previews."),
        ("--force", "Off", "Override rebalance guard (skip schedule/drift checks)"),
        ("--status", "Off", "Only sync equity and positions, don't generate orders"),
        ("--sync", "Off", "Reconcile submitted orders against Moomoo order history"),
        ("--allow-closed-market-submit", "Off", "Allow orders outside market hours (9:30-16:00 ET)"),
        ("--allow-stale-signal", "Off", "Submit even if signal/factor data is old"),
        ("--etf-drift-threshold", "0.03", "Minimum ETF weight drift before trading (3%)"),
        ("--overlay-drift-threshold", "0.01", "Minimum stock weight drift before trading (1%)"),
        ("--min-trade-value", "$25", "Skip orders smaller than this dollar amount"),
        ("--limit-offset-bps", "10", "Limit price offset in basis points (0.10%)"),
        ("--max-signal-age-hours", "24", "Block trading if signal is older than this"),
        ("--max-factor-age-trading-days", "5", "Block trading if factor data is older than this"),
        ("--equity", "Auto", "Override account equity for position sizing"),
        ("--max-submit-gross-exposure", "1.00", "Maximum total exposure submitted (1.0 = no leverage)"),
        ("--reset-paper-log", "Off", "Archive old trade log before this submission"),
    ]
    for p in mparams:
        add_table_row(moomoo_params, p)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 6: alpaca_paper_trading.py
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "6. alpaca_paper_trading.py - Alpaca Broker Script", level=1)

    add_para(doc, "Purpose", bold=True)
    add_para(doc, (
        "This script reads the TQQQ-enhanced signal and submits orders to Alpaca's free "
        "paper trading API. Unlike Moomoo (which requires a desktop app), Alpaca works "
        "entirely through a cloud API - no desktop software needed."
    ))

    add_para(doc, "Prerequisites", bold=True)
    add_bullet(doc, "Sign up at https://alpaca.markets (free, instant)")
    add_bullet(doc, "pip install alpaca-trade-api")
    add_bullet(doc, "Set ALPACA_API_KEY and ALPACA_SECRET_KEY in your .env file")

    add_para(doc, "How to Run", bold=True)
    add_code(doc, "  python3 alpaca_paper_trading.py                  # Preview orders (dry run)")
    add_code(doc, "  python3 alpaca_paper_trading.py --submit         # Submit orders + auto-snapshot equity")
    add_code(doc, "  python3 alpaca_paper_trading.py --status         # Show account + positions")
    add_code(doc, "  python3 alpaca_paper_trading.py --reconcile      # Check if pending orders filled")
    add_code(doc, "  python3 alpaca_paper_trading.py --force          # Skip drift + duplicate checks")

    add_para(doc, "Key Logic Explained", bold=True)

    add_para(doc, "1. AlpacaBroker Class", italic=True)
    add_para(doc, (
        "Implements the Broker interface. Methods include:\n"
        "- get_equity(): Total account value (cash + positions)\n"
        "- get_cash(): Available cash balance\n"
        "- get_positions(): List of all current holdings\n"
        "- place_order(): Submit a buy or sell order\n"
        "- is_market_open(): Check if US market is currently open\n"
        "- get_last_price(): Get latest price with retry logic"
    ))

    add_para(doc, "2. Retry Logic on Price Fetches (get_last_price)", italic=True)
    add_para(doc, (
        "When fetching a stock price, network errors or API throttling can cause failures. "
        "Instead of silently returning 0 and skipping the ticker, the system:\n"
        "- Tries the primary endpoint (get_latest_trade)\n"
        "- Falls back to the snapshot endpoint\n"
        "- Retries up to 3 times with exponential backoff (1s, 2s, 4s delays)\n"
        "- Logs every failure so you can diagnose persistent issues\n"
        "- Only returns 0.0 after all retries are exhausted"
    ))

    add_para(doc, "3. Duplicate Order Prevention (_already_submitted_today)", italic=True)
    add_para(doc, (
        "If you accidentally run --submit twice in one day, the second run will be blocked "
        "with a warning message. It checks the trade log (alpaca_paper_log.csv) for any "
        "submissions from today's date. Use --force to override this check."
    ))

    add_para(doc, "4. parse_target_weights(signal)", italic=True)
    add_para(doc, (
        "Extracts target portfolio weights from the signal CSV. Handles:\n"
        "- target_spy_weight (SPY allocation)\n"
        "- target_qqq_weight (QQQ allocation)\n"
        "- target_tqqq_weight (TQQQ allocation, unique to this strategy)\n"
        "- overlay_weights_json (individual stock picks as JSON)\n"
        "Weights of 0.0 are excluded. Malformed JSON degrades gracefully to empty overlay."
    ))

    add_para(doc, "5. scale_weights(weights, max_gross)", italic=True)
    add_para(doc, (
        "If total gross exposure exceeds the limit (default 1.0), all weights are "
        "multiplied by the same scale factor. Example: if weights sum to 1.25 and "
        "max is 1.0, every weight is multiplied by 0.8 (= 1.0/1.25)."
    ))

    add_para(doc, "6. generate_orders(broker, target_weights, force)", italic=True)
    add_para(doc, (
        "Compares target weights to current positions and generates orders:\n"
        "a) Fetches current equity and position quantities\n"
        "b) For each ticker, calculates current weight vs target weight\n"
        "c) Computes 'drift' (the difference)\n"
        "d) Only generates an order if drift exceeds the threshold (3% for ETFs, 1% for stocks)\n"
        "e) Sorts orders: sells first (free up cash), then buys by size\n"
        "f) Skips orders smaller than $25 (the MIN_TRADE_VALUE)"
    ))

    add_para(doc, "7. Auto Equity Snapshot (snapshot_equity)", italic=True)
    add_para(doc, (
        "After every --submit, the script automatically saves today's account equity "
        "to alpaca_paper_equity.csv. This builds the equity history that the gauntlet "
        "needs to compute Sharpe ratio and drawdown. Same-day duplicates are replaced "
        "with the latest value."
    ))

    add_para(doc, "8. Order Reconciliation (reconcile_orders)", italic=True)
    add_para(doc, (
        "Queries Alpaca's API for the real status of each submitted order. Updates "
        "the log with: filled, cancelled, partially_filled, etc. This is important "
        "for the gauntlet's fill rate calculation."
    ))

    add_para(doc, "Key Parameters", bold=True)
    alpaca_params = doc.add_table(rows=1, cols=3)
    alpaca_params.style = "Light Grid Accent 1"
    hdr = alpaca_params.rows[0].cells
    hdr[0].text = "Parameter"
    hdr[1].text = "Default"
    hdr[2].text = "What It Does"

    aparams = [
        ("--submit", "Off", "Actually submit orders to Alpaca"),
        ("--status", "Off", "Just show account status, no orders"),
        ("--reconcile", "Off", "Check if pending orders filled"),
        ("--force", "Off", "Skip drift thresholds AND duplicate-day check"),
        ("ALPACA_API_KEY", "(env var)", "Your Alpaca paper trading API key"),
        ("ALPACA_SECRET_KEY", "(env var)", "Your Alpaca paper trading secret"),
        ("ALPACA_ETF_DRIFT_THRESHOLD", "0.03", "Minimum ETF drift before trading"),
        ("ALPACA_OVERLAY_DRIFT_THRESHOLD", "0.01", "Minimum stock drift before trading"),
        ("ALPACA_MIN_TRADE_VALUE", "$25", "Skip orders smaller than this"),
        ("ALPACA_MAX_GROSS_EXPOSURE", "1.00", "Maximum total portfolio exposure"),
        ("ALPACA_PRICE_RETRIES", "3", "Number of price fetch retry attempts"),
        ("ALPACA_PRICE_RETRY_DELAY", "1.0", "Base delay in seconds between retries"),
    ]
    for p in aparams:
        add_table_row(alpaca_params, p)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 7: paper_gauntlet.py
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "7. paper_gauntlet.py - Moomoo Health Check", level=1)

    add_para(doc, "Purpose", bold=True)
    add_para(doc, (
        "The paper gauntlet is a 'final exam' for the strategy. Before you put real money "
        "into this system, you need proof it works in live markets. This script checks "
        "every aspect of the Moomoo paper trading performance and gives a pass/fail verdict."
    ))

    add_para(doc, "How to Run", bold=True)
    add_code(doc, "  python3 paper_gauntlet.py")

    add_para(doc, "What It Checks", bold=True)

    gauntlet_table = doc.add_table(rows=1, cols=4)
    gauntlet_table.style = "Light Grid Accent 1"
    hdr = gauntlet_table.rows[0].cells
    hdr[0].text = "Check"
    hdr[1].text = "Threshold"
    hdr[2].text = "Why It Matters"
    hdr[3].text = "Env Var Override"

    checks = [
        ("Trading Days", ">= 20 days", "Need enough data for statistical significance", "PAPER_GAUNTLET_MIN_EQUITY_DAYS"),
        ("Sharpe Ratio", ">= 0.5", "Risk-adjusted return must justify the effort", "N/A (hardcoded)"),
        ("Max Drawdown", ">= -10%", "Account can't drop more than 10% from peak", "N/A (hardcoded)"),
        ("Fill Rate", ">= 95%", "Orders must actually execute at the broker", "PAPER_GAUNTLET_MIN_FILL_RATE"),
        ("Cancel Rate", "<= 5%", "Too many cancelled orders means something is wrong", "PAPER_GAUNTLET_MAX_CANCEL_RATE"),
        ("Portfolio Drift", "<= 15%", "Current positions must be close to targets", "PAPER_GAUNTLET_MAX_DRIFT"),
        ("Avg Slippage", "<= 10 bps", "Execution costs must be reasonable", "PAPER_GAUNTLET_MAX_AVG_SLIPPAGE_BPS"),
        ("Signal Freshness", "paper_ready=True", "Strategy must be approved by backtest gates", "N/A"),
        ("Survivorship Stress", "Pass", "Strategy must survive removing stocks that failed audit", "N/A"),
        ("Execution Stress", "Pass", "Strategy must survive adverse execution scenarios", "N/A"),
        ("Factor Decay", "No block", "Factor model must not show signs of decay", "N/A"),
    ]
    for c in checks:
        add_table_row(gauntlet_table, c)

    add_para(doc, "")
    add_para(doc, "Key Logic: evaluate_paper()", italic=True)
    add_para(doc, (
        "This is the main evaluation function. It:\n"
        "1. Reads paper_trades.csv (order history), paper_equity.csv (daily equity), "
        "and paper_daily_status.json (latest positions/regime)\n"
        "2. Computes performance metrics: total return, Sharpe ratio, max drawdown\n"
        "3. Checks order fill quality: how many orders actually filled vs cancelled\n"
        "4. Validates signal freshness and factor model health\n"
        "5. Compares paper returns against SPY/QQQ benchmarks over the same period\n"
        "6. Checks additional stress tests (survivorship, execution)\n"
        "7. Returns a pass/fail verdict with detailed reasons for any failure"
    ))

    add_para(doc, "Output", bold=True)
    add_bullet(doc, "logs/paper_gauntlet_YYYYMMDD.json - Full evaluation report")
    add_bullet(doc, "Console output shows pass/fail status and reason")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 8: alpaca_paper_gauntlet.py
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "8. alpaca_paper_gauntlet.py - Alpaca Health Check", level=1)

    add_para(doc, "Purpose", bold=True)
    add_para(doc, (
        "Same concept as the Moomoo gauntlet, but for the Alpaca TQQQ strategy. The key "
        "difference is a wider max drawdown threshold (-15% vs -10%) because TQQQ's 3x "
        "leverage causes bigger swings."
    ))

    add_para(doc, "How to Run", bold=True)
    add_code(doc, "  python3 alpaca_paper_gauntlet.py              # Full evaluation")
    add_code(doc, "  python3 alpaca_paper_gauntlet.py --verbose    # Extra detail")
    add_code(doc, "  python3 alpaca_paper_gauntlet.py --snapshot   # Save today's equity")
    add_code(doc, "  python3 alpaca_paper_gauntlet.py --json       # Raw JSON output")

    add_para(doc, "What It Checks", bold=True)

    alpaca_gauntlet_table = doc.add_table(rows=1, cols=3)
    alpaca_gauntlet_table.style = "Light Grid Accent 1"
    hdr = alpaca_gauntlet_table.rows[0].cells
    hdr[0].text = "Gate"
    hdr[1].text = "Threshold"
    hdr[2].text = "Description"

    alpaca_checks = [
        ("Trading Days", ">= 20", "Enough data for meaningful statistics"),
        ("Fill Rate", ">= 95%", "Orders must actually fill at Alpaca"),
        ("Cancel Rate", "<= 5%", "Too many cancels means order problems"),
        ("Signal Freshness", "<= 48 hours old", "Signal must be recent"),
        ("Sharpe Ratio", ">= 0.5", "Risk-adjusted return must be adequate"),
        ("Max Drawdown", ">= -15%", "Wider than Moomoo because TQQQ is more volatile"),
        ("Portfolio Drift", "<= 15%", "Positions must track signal targets"),
        ("Alpaca Connection", "Must connect", "Live connection to verify account state"),
    ]
    for c in alpaca_checks:
        add_table_row(alpaca_gauntlet_table, c)

    add_para(doc, "")
    add_para(doc, "Key Difference from Moomoo Gauntlet", bold=True)
    add_para(doc, (
        "The Alpaca gauntlet also checks the live account state by connecting to Alpaca's "
        "API. It pulls current equity, positions, and computes real-time drift between "
        "what you hold and what the signal says you should hold. The Moomoo gauntlet "
        "relies on saved status files instead."
    ))

    add_para(doc, "Parameters", bold=True)
    alpaca_g_params = doc.add_table(rows=1, cols=3)
    alpaca_g_params.style = "Light Grid Accent 1"
    hdr = alpaca_g_params.rows[0].cells
    hdr[0].text = "Parameter"
    hdr[1].text = "Default"
    hdr[2].text = "What It Does"

    agparams = [
        ("--verbose", "Off", "Show extra detail in report"),
        ("--snapshot", "Off", "Save today's equity to tracking CSV"),
        ("--json", "Off", "Output raw JSON instead of formatted report"),
        ("ALPACA_GAUNTLET_MIN_DAYS", "20", "Minimum trading days required"),
        ("ALPACA_GAUNTLET_MIN_SHARPE", "0.5", "Minimum Sharpe ratio"),
        ("ALPACA_GAUNTLET_MAX_DD", "-15.0", "Max allowed drawdown (%)"),
        ("ALPACA_GAUNTLET_MIN_FILL_RATE", "0.95", "Minimum order fill rate"),
        ("ALPACA_GAUNTLET_MAX_SIGNAL_AGE_HOURS", "48", "Max signal age in hours"),
    ]
    for p in agparams:
        add_table_row(alpaca_g_params, p)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 9: daily_run.py
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "9. daily_run.py - Automation Runner", level=1)

    add_para(doc, "Purpose", bold=True)
    add_para(doc, (
        "Instead of running 8 separate commands every trading day, this script chains "
        "them all together in one command. If one step fails, it logs the error and "
        "continues with the rest - a bad signal generation won't prevent the other "
        "strategy from trading."
    ))

    add_para(doc, "How to Run", bold=True)
    add_code(doc, "  python3 daily_run.py              # Run everything")
    add_code(doc, "  python3 daily_run.py --dry-run    # Show what would run without executing")
    add_code(doc, "  python3 daily_run.py --moomoo     # Only run Moomoo steps")
    add_code(doc, "  python3 daily_run.py --alpaca     # Only run Alpaca steps")
    add_code(doc, "  python3 daily_run.py --timeout 600  # 10 min timeout per step (default 5 min)")

    add_para(doc, "Steps Executed (in order)", bold=True)
    step_table = doc.add_table(rows=1, cols=4)
    step_table.style = "Light Grid Accent 1"
    hdr = step_table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Step Name"
    hdr[2].text = "Command"
    hdr[3].text = "What It Does"

    daily_steps = [
        ("1", "moomoo_signal", "core_satellite_alpha.py", "Generate Moomoo signal"),
        ("2", "moomoo_submit", "moomoo_paper_trading.py --submit", "Submit orders to Moomoo"),
        ("3", "moomoo_status", "moomoo_paper_trading.py --status", "Sync equity and positions"),
        ("4", "moomoo_gauntlet", "paper_gauntlet.py", "Health check for Moomoo"),
        ("5", "alpaca_signal", "core_satellite_tqqq.py", "Generate TQQQ signal"),
        ("6", "alpaca_submit", "alpaca_paper_trading.py --submit", "Submit orders to Alpaca"),
        ("7", "alpaca_reconcile", "alpaca_paper_trading.py --reconcile", "Check if Alpaca orders filled"),
        ("8", "alpaca_gauntlet", "alpaca_paper_gauntlet.py", "Health check for Alpaca"),
    ]
    for s in daily_steps:
        add_table_row(step_table, s)

    add_para(doc, "")
    add_para(doc, "Key Logic: run_step()", italic=True)
    add_para(doc, (
        "Each step runs as a subprocess with a 5-minute timeout (configurable). The function:\n"
        "- Prints the step name and command being run\n"
        "- Captures stdout and stderr\n"
        "- Shows the last 20 lines of output for readability\n"
        "- Records the result (ok/failed/timeout/error) with elapsed time\n"
        "- A failure does NOT stop subsequent steps"
    ))

    add_para(doc, "Scheduling with Cron", bold=True)
    add_para(doc, (
        "To run automatically every weekday at market open:"
    ))
    add_code(doc, '  30 9 * * 1-5 cd "/path/to/Stock Market AI Bot" && python3 daily_run.py >> logs/daily_run.log 2>&1')

    add_para(doc, "Output", bold=True)
    add_bullet(doc, "Console: Summary of passed/failed/skipped steps with total time")
    add_bullet(doc, "logs/daily_run_YYYYMMDD.json: Full run log with all step results")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 10: paper_report.py
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "10. paper_report.py - Performance Comparison", level=1)

    add_para(doc, "Purpose", bold=True)
    add_para(doc, (
        "A report card that shows how both strategies are performing side-by-side. "
        "Compares Moomoo (core-satellite) vs Alpaca (TQQQ-enhanced) across return, "
        "Sharpe, drawdown, fill rate, and benchmarks."
    ))

    add_para(doc, "How to Run", bold=True)
    add_code(doc, "  python3 paper_report.py           # Formatted report")
    add_code(doc, "  python3 paper_report.py --json    # Raw JSON output")

    add_para(doc, "What It Shows", bold=True)
    add_bullet(doc, "Data Coverage: How many equity snapshots and trading days each strategy has")
    add_bullet(doc, "Account: Starting vs current equity, gross exposure")
    add_bullet(doc, "Performance: Total return, Sharpe ratio, max drawdown")
    add_bullet(doc, "Order Execution: Total orders, filled, cancelled, fill rate")
    add_bullet(doc, "Benchmarks: SPY and QQQ returns over the same period")
    add_bullet(doc, "Alpha: How much each strategy beat (or trailed) SPY/QQQ")
    add_bullet(doc, "Verdict: Which strategy is currently winning")

    add_para(doc, "Data Sources", bold=True)
    data_table = doc.add_table(rows=1, cols=3)
    data_table.style = "Light Grid Accent 1"
    hdr = data_table.rows[0].cells
    hdr[0].text = "File"
    hdr[1].text = "Strategy"
    hdr[2].text = "Contents"

    data_sources = [
        ("paper_equity.csv", "Moomoo", "Daily equity snapshots (date, equity, regime)"),
        ("paper_trades.csv", "Moomoo", "Order submission log with fill status"),
        ("paper_daily_status.json", "Moomoo", "Latest positions, drift, regime"),
        ("alpaca_paper_equity.csv", "Alpaca", "Daily equity snapshots (date, equity, cash)"),
        ("alpaca_paper_log.csv", "Alpaca", "Order submission log with fill status"),
        ("core_satellite_tqqq_signal.csv", "Alpaca", "Latest TQQQ signal (for regime info)"),
    ]
    for d in data_sources:
        add_table_row(data_table, d)

    add_para(doc, "")
    add_para(doc, "Output", bold=True)
    add_bullet(doc, "Console: Formatted side-by-side comparison table")
    add_bullet(doc, "logs/paper_report_YYYYMMDD.json: Full report data saved for later")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 11: settings.py
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "11. settings.py - Configuration", level=1)

    add_para(doc, "Purpose", bold=True)
    add_para(doc, (
        "Central configuration file for the entire project. All other scripts import "
        "their settings from here. Uses a .env file for sensitive values (API keys)."
    ))

    add_para(doc, "Key Settings Used by Paper Trading", bold=True)
    settings_table = doc.add_table(rows=1, cols=3)
    settings_table.style = "Light Grid Accent 1"
    hdr = settings_table.rows[0].cells
    hdr[0].text = "Setting"
    hdr[1].text = "Default"
    hdr[2].text = "Description"

    settings = [
        ("SIGNAL_DIR", "'signals'", "Directory where signal CSVs and trade logs are stored"),
        ("LOG_DIR", "'logs'", "Directory for gauntlet reports and run logs"),
        ("DATA_DIR", "'data'", "Directory with stock price parquet files"),
        ("PAPER_MODE_STRATEGY", "'core_satellite_alpha'", "Which strategy the Moomoo script uses"),
        ("SLIPPAGE_BASE_PCT", "0.001", "Estimated transaction cost per trade (0.1%)"),
        ("SINGLE_NAME_PAPER_TRADING_ENABLED", "False", "Whether individual stock trading is enabled"),
        ("FINNHUB_API_KEY", "(from .env)", "API key for Finnhub data provider"),
        ("ALPACA_API_KEY", "(from .env)", "API key for Alpaca paper trading"),
        ("ALPACA_SECRET_KEY", "(from .env)", "Secret key for Alpaca paper trading"),
    ]
    for s in settings:
        add_table_row(settings_table, s)

    add_para(doc, "")
    add_para(doc, "The .env File", bold=True)
    add_para(doc, (
        "Sensitive values like API keys are stored in a .env file in the project root. "
        "This file is NOT committed to version control (it's in .gitignore). Example:"
    ))
    add_code(doc, (
        "  FINNHUB_API_KEY=your_finnhub_key\n"
        "  ALPACA_API_KEY=your_alpaca_key\n"
        "  ALPACA_SECRET_KEY=your_alpaca_secret\n"
        "  PAPER_MODE_STRATEGY=core_satellite_alpha\n"
        "  PAPER_SIGNAL_TIMEZONE=Asia/Singapore"
    ))

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 12: GLOSSARY
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "12. Glossary of Key Terms", level=1)

    terms = [
        ("Alpha", "The extra return your strategy earns above a benchmark (like SPY). If SPY returns 10% and you return 15%, your alpha is +5%."),
        ("Backtest", "Testing a trading strategy on historical data to see how it would have performed in the past."),
        ("Basis Points (bps)", "1/100th of a percent. 100 bps = 1%. Used to measure tiny price differences."),
        ("Benchmark", "A standard to compare your strategy against. Usually SPY (S&P 500) or QQQ (Nasdaq 100)."),
        ("CAGR", "Compound Annual Growth Rate. The average yearly return if gains were reinvested."),
        ("Core-Satellite", "Investment approach: safe ETFs (core) + stock picks (satellite)."),
        ("Drawdown", "The percentage drop from your account's highest point to its lowest point. -10% means you lost 10% from peak."),
        ("Drift", "The difference between what you currently hold and what the strategy wants you to hold."),
        ("ETF", "Exchange-Traded Fund. A basket of stocks you can buy as a single ticker (SPY, QQQ, TQQQ)."),
        ("Factor Score", "A number measuring how attractive a stock is based on quantitative metrics like momentum, value, and quality."),
        ("Fill Rate", "The percentage of submitted orders that actually executed at the broker."),
        ("Gauntlet", "A series of health checks the strategy must pass before real money is deployed."),
        ("Gross Exposure", "Total invested amount as a fraction of account equity. 1.0 = 100% invested, 1.25 = using 25% leverage."),
        ("Leverage", "Investing more money than you have by borrowing. TQQQ uses 3x leverage internally."),
        ("Moving Average", "The average price over a set number of days. If the current price is above the moving average, it's in an uptrend."),
        ("Overlay", "The 'satellite' part of core-satellite. Individual stock picks layered on top of the ETF core."),
        ("Paper Trading", "Simulated trading with fake money to test strategies before using real capital."),
        ("Rebalance", "Adjusting your portfolio by buying/selling to match the strategy's target weights."),
        ("Regime", "The current market condition: risk_on (bullish), neutral, or risk_off (bearish)."),
        ("Sharpe Ratio", "Return divided by risk (volatility). Higher is better. Above 1.0 is good, above 2.0 is excellent."),
        ("Slippage", "The difference between the price you expected to trade at and the price you actually got."),
        ("Sticky Scoring", "A method that blends new weights with old ones (65% old, 35% new) to reduce unnecessary trading."),
        ("TQQQ", "ProShares UltraPro QQQ. A 3x leveraged ETF that triples QQQ's daily returns (and losses)."),
        ("Volatility", "How much a price moves up and down. High volatility = big swings."),
        ("Walkforward", "A backtest method where the model is trained on past data and tested on future data, rolling forward through time."),
    ]

    for term, definition in terms:
        add_para(doc, term, bold=True)
        add_para(doc, definition)

    return doc


if __name__ == "__main__":
    doc = build_document()
    out_dir = Path("Documentation")
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / "doc_core_paper_trading_system.docx"
    doc.save(str(out_path))
    print(f"Documentation saved -> {out_path}")
